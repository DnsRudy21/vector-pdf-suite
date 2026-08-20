import io, json, re, zipfile
from pathlib import Path
from typing import Callable
import fitz
from docx import Document
from openpyxl import Workbook
from PIL import Image

Progress = Callable[[int], None]
def safe_stem(name: str) -> str:
    return (re.sub(r"[^A-Za-z0-9._-]+", "-", Path(name).stem).strip(".-")[:80] or "document")

class PDFConverter:
    IMAGE_FORMATS = {"png":"PNG", "jpg":"JPEG", "jpeg":"JPEG", "bmp":"BMP", "tiff":"TIFF"}

    @staticmethod
    def validate_pdf(source: Path) -> int:
        try:
            with fitz.open(source) as doc:
                if not doc.is_pdf or doc.page_count < 1: raise ValueError("El archivo no es un PDF válido")
                return doc.page_count
        except (fitz.FileDataError, fitz.EmptyFileError) as exc: raise ValueError("PDF inválido o dañado") from exc

    @classmethod
    def pdf_to_images(cls, sources: list[Path], output: Path, options: dict, progress: Progress) -> None:
        key = str(options.get("format","png")).lower(); fmt = cls.IMAGE_FORMATS.get(key)
        dpi, quality = int(options.get("dpi",150)), int(options.get("quality",85))
        if not fmt or not 72 <= dpi <= 600 or not 1 <= quality <= 100: raise ValueError("Parámetros de imagen inválidos")
        total = sum(cls.validate_pdf(p) for p in sources); done = 0
        with zipfile.ZipFile(output,"w",zipfile.ZIP_DEFLATED) as archive:
            for source in sources:
                with fitz.open(source) as doc:
                    for index,page in enumerate(doc):
                        pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72,dpi/72),alpha=False)
                        image = Image.open(io.BytesIO(pix.tobytes("png"))); data = io.BytesIO()
                        image.save(data,fmt,**({"quality":quality,"optimize":True} if fmt=="JPEG" else {}))
                        archive.writestr(f"{safe_stem(source.name)}/page-{index+1:04d}.{key}",data.getvalue())
                        done += 1; progress(round(done/total*95))

    @staticmethod
    def images_to_pdf(sources: list[Path], output: Path, options: dict, progress: Progress) -> None:
        images=[]
        try:
            for index,source in enumerate(sources):
                with Image.open(source) as original:
                    image=original.convert("RGB"); image.thumbnail((4961,7016),Image.Resampling.LANCZOS); images.append(image.copy())
                progress(round((index+1)/len(sources)*70))
            if not images: raise ValueError("No se recibieron imágenes")
            images[0].save(output,"PDF",save_all=True,append_images=images[1:],resolution=float(options.get("dpi",150)),quality=int(options.get("quality",85)))
        finally:
            for image in images: image.close()

    @classmethod
    def merge(cls,sources:list[Path],output:Path,_options:dict,progress:Progress)->None:
        result=fitz.open()
        try:
            for index,source in enumerate(sources):
                cls.validate_pdf(source)
                with fitz.open(source) as doc: result.insert_pdf(doc)
                progress(round((index+1)/len(sources)*90))
            result.save(output,garbage=4,deflate=True)
        finally: result.close()

    @classmethod
    def split(cls,sources:list[Path],output:Path,options:dict,progress:Progress)->None:
        with zipfile.ZipFile(output,"w",zipfile.ZIP_DEFLATED) as archive:
            for file_index,source in enumerate(sources):
                with fitz.open(source) as doc:
                    groups=[[i] for i in range(doc.page_count)] if options.get("split_mode","each")=="each" else cls._ranges(str(options.get("ranges","")),doc.page_count)
                    for group_index,pages in enumerate(groups):
                        part=fitz.open()
                        for page in pages: part.insert_pdf(doc,from_page=page,to_page=page)
                        archive.writestr(f"{safe_stem(source.name)}-parte-{group_index+1}.pdf",part.tobytes(garbage=4,deflate=True)); part.close()
                progress(round((file_index+1)/len(sources)*95))

    @staticmethod
    def _ranges(raw:str,count:int)->list[list[int]]:
        if not raw: raise ValueError("Indica rangos como 1-3,4,5-7")
        groups=[]
        for token in raw.split(","):
            parts=token.strip().split("-"); start=int(parts[0]); end=int(parts[-1])
            if start<1 or end<start or end>count: raise ValueError(f"Rango inválido: {token}")
            groups.append(list(range(start-1,end)))
        return groups

    @classmethod
    def _batch_pdf(cls,sources:list[Path],output:Path,progress:Progress,action:Callable)->None:
        with zipfile.ZipFile(output,"w",zipfile.ZIP_DEFLATED) as archive:
            for index,source in enumerate(sources):
                cls.validate_pdf(source); temp=output.parent/f"result-{index}.pdf"
                with fitz.open(source) as doc: action(doc,temp)
                archive.write(temp,f"{safe_stem(source.name)}-procesado.pdf"); temp.unlink(missing_ok=True); progress(round((index+1)/len(sources)*95))

    @classmethod
    def optimize(cls,sources:list[Path],output:Path,_options:dict,progress:Progress)->None:
        cls._batch_pdf(sources,output,progress,lambda doc,target:doc.save(target,garbage=4,clean=True,deflate=True,deflate_images=True,deflate_fonts=True))

    @classmethod
    def protect(cls,sources:list[Path],output:Path,options:dict,progress:Progress)->None:
        password=str(options.get("password",""))
        if len(password)<4: raise ValueError("La contraseña debe tener al menos 4 caracteres")
        perms=fitz.PDF_PERM_PRINT|(fitz.PDF_PERM_COPY if options.get("allow_copy",False) else 0)
        cls._batch_pdf(sources,output,progress,lambda doc,target:doc.save(target,encryption=fitz.PDF_ENCRYPT_AES_256,owner_pw=password,user_pw=password,permissions=perms,garbage=4,deflate=True))

    @classmethod
    def watermark(cls,sources:list[Path],output:Path,options:dict,progress:Progress)->None:
        text=str(options.get("watermark","CONFIDENCIAL"))[:120]; opacity=max(.05,min(float(options.get("opacity",.2)),1))
        def save(doc,target):
            for page in doc:
                point=fitz.Point(page.rect.width*.12,page.rect.height*.52); page.insert_text(point,text,fontsize=max(20,page.rect.width/14),color=(.5,.5,.5),fill_opacity=opacity)
            doc.save(target,garbage=4,deflate=True)
        cls._batch_pdf(sources,output,progress,save)

    @classmethod
    def extract(cls,sources:list[Path],output:Path,options:dict,progress:Progress)->None:
        target_type=options.get("target","word")
        with zipfile.ZipFile(output,"w",zipfile.ZIP_DEFLATED) as archive:
            for index,source in enumerate(sources):
                with fitz.open(source) as pdf: pages=[page.get_text("text") for page in pdf]
                if target_type=="excel":
                    book=Workbook(); sheet=book.active; sheet.append(["Página","Texto"])
                    for number,text in enumerate(pages,1): sheet.append([number,text])
                    temp=output.parent/f"extract-{index}.xlsx"; book.save(temp); ext="xlsx"
                else:
                    doc=Document(); doc.add_heading(source.name,0)
                    for number,text in enumerate(pages,1): doc.add_heading(f"Página {number}",1); doc.add_paragraph(text)
                    temp=output.parent/f"extract-{index}.docx"; doc.save(temp); ext="docx"
                archive.write(temp,f"{safe_stem(source.name)}.{ext}"); temp.unlink(missing_ok=True); progress(round((index+1)/len(sources)*95))

    @classmethod
    def metadata(cls,sources:list[Path],output:Path,_options:dict,progress:Progress)->None:
        records=[]
        for index,source in enumerate(sources):
            with fitz.open(source) as pdf: records.append({"file":source.name,"pages":pdf.page_count,**pdf.metadata})
            progress(round((index+1)/len(sources)*95))
        output.write_text(json.dumps(records,ensure_ascii=False,indent=2),encoding="utf-8")

    OPERATIONS={"pdf-to-images":pdf_to_images,"images-to-pdf":images_to_pdf,"merge":merge,"split":split,"compress":optimize,"protect":protect,"watermark":watermark,"extract":extract,"metadata":metadata}
